# -*- coding: utf-8 -*-
"""生成《法言白话》初赛作品说明文档（技术方案按参考方案目录扩写）"""
from __future__ import annotations

import os
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
OUT = os.path.join(ROOT, 'doc', '01-作品说明文档-法言白话队.docx')
ASSETS = os.path.join(ROOT, 'doc', 'assets')

import sys
sys.path.insert(0, os.path.dirname(__file__))
from section_43 import write_section_43  # noqa: E402


def set_run_font(run, name='宋体', size=Pt(10.5), bold=False, color=None):
    run.bold = bold
    run.font.size = size
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if color is not None:
        run.font.color.rgb = color


def set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=None, space_before=0, space_after=3):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    p.alignment = align
    if first_line is not None:
        pf.first_line_indent = first_line


class Builder:
    def __init__(self):
        self.doc = Document()
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)
        self._add_page_number(section)

    def _add_page_number(self, section):
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        set_run_font(run, size=Pt(10.5))
        fld1 = OxmlElement('w:fldChar')
        fld1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        fld2 = OxmlElement('w:fldChar')
        fld2.set(qn('w:fldCharType'), 'end')
        run._r.append(fld1)
        run._r.append(instr)
        run._r.append(fld2)

    def text(self, text, size=Pt(10.5), bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             first_line=None, space_before=0, space_after=3, color=None):
        p = self.doc.add_paragraph()
        set_para_format(p, align=align, first_line=first_line, space_before=space_before, space_after=space_after)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        return p

    def h(self, text, level=1):
        sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(10.5)}
        return self.text(text, size=sizes.get(level, Pt(10.5)), bold=True, space_before=8, space_after=4)

    def body(self, text, space_after=3):
        return self.text(text, first_line=Cm(0.74), space_after=space_after)

    def bullet(self, text, space_after=2):
        return self.text('• ' + text, space_after=space_after)

    def caption(self, text):
        return self.text(text, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER,
                         color=RGBColor(0x55, 0x55, 0x55), space_after=6)

    def picture(self, name, width_cm=13.5):
        path = os.path.join(ASSETS, name)
        if os.path.exists(path):
            p = self.doc.add_paragraph()
            set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)
            run = p.add_run()
            run.add_picture(path, width=Cm(width_cm))
            return p
        return None

    def shade(self, cell, fill='D9E2F3'):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def cell(self, cell, text, bold=False, size=Pt(10.5), center=False):
        cell.text = ''
        p = cell.paragraphs[0]
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT, space_after=0)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def table(self, rows):
        t = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                self.cell(t.rows[i].cells[j], val, bold=(i == 0), size=Pt(9), center=(i == 0 or j == 0))
                if i == 0:
                    self.shade(t.rows[i].cells[j])
        return t

    def page_break(self):
        self.doc.add_page_break()

    def save(self):
        os.makedirs(os.path.dirname(OUT), exist_ok=True)
        self.doc.save(OUT)
        print('saved', OUT)


def build():
    b = Builder()

    # ========== 封面 ==========
    b.text('2026中国高校计算机大赛—人工智能创意赛', size=Pt(22), bold=True,
           align=WD_ALIGN_PARAGRAPH.CENTER, space_before=28, space_after=6)
    b.text('鸿蒙赛道作品说明文档（初赛）', size=Pt(22), bold=True,
           align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    for label, val in [
        ('参赛学校：', '【待填写：以队长学校为准】'),
        ('团队名称：', '法言白话队'),
        ('作品名称：', '法言白话'),
        ('赛题方向：', '应用创新'),
        ('联系人（队长）：', '【待填写】'),
        ('联系电话（队长）：', '【待填写】'),
    ]:
        p = b.doc.add_paragraph()
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        r1 = p.add_run(label)
        set_run_font(r1, size=Pt(14), bold=True)
        r2 = p.add_run(val)
        set_run_font(r2, size=Pt(14))
    b.page_break()

    # ========== 目录 ==========
    b.text('作品说明文档（初赛）', size=Pt(18), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    b.text('包含但不限于以下内容：', space_after=3)
    for item in [
        '一、参赛团队信息表',
        '二、作品原创性声明',
        '三、创意描述（30字以内）',
        '四、设计稿 / 技术方案（含技术路线与分模块实现）',
        '五、介绍文档（800字以内）',
    ]:
        b.text(item, space_after=4)

    b.h('技术方案章节目录（参考项目详细方案体例）', 2)
    for item in [
        '4.1 产品定位与赛题契合',
        '4.2 交互设计稿',
        '4.3 技术路线（文字介绍为主）',
        '　　4.3.1 总体思路与设计原则',
        '　　4.3.2 系统体系架构',
        '　　4.3.3 端到端业务数据流',
        '　　4.3.4 技术框架选型说明',
        '　　4.3.5 客户端技术路线',
        '　　4.3.6 分析服务技术路线',
        '　　4.3.7 合同追问技术路线',
        '　　4.3.8 知识库技术路线',
        '　　4.3.9 鸿蒙全场景能力路线',
        '　　4.3.10 稳定性与演示兜底',
        '　　4.3.11 运行与协作方式',
        '　　4.3.12 边界说明',
        '　　4.3.13 小结',
        '4.4 技术实现方案',
        '　　4.4.1～4.4.9 按业务模块做文字说明（导入、分析、知识、追问、状态、协作、界面、鸿蒙、导出兜底）',
        '4.5 系统实现效果',
        '4.6 运行与使用说明',
    ]:
        b.text(item, space_after=2)
    # 目录后紧接团队信息，避免目录页大块留白
    b.h('一、参赛团队信息表', 1)
    info = b.doc.add_table(rows=5, cols=2)
    info.style = 'Table Grid'
    for i, (k, v) in enumerate([
        ('作品名称', '法言白话（30个字符以内）'),
        ('团队名称', '法言白话队（10个字符以内，无标点）'),
        ('参赛学校', '【待填写】（以队长学校为准）'),
        ('赛题方向', '（√）应用创新   （ ）Agent创新   （ ）用户体验创新   （ ）操作系统智能创新'),
        ('作品定位', '面向非法律专业用户的合同智能解读鸿蒙应用'),
    ]):
        b.cell(info.rows[i].cells[0], k, bold=True, center=True)
        b.shade(info.rows[i].cells[0])
        b.cell(info.rows[i].cells[1], v)

    b.text('团队队员基本信息（可跨校、跨专业组队）', size=Pt(12), bold=True, space_before=8, space_after=4)
    headers = ['姓名', '学校全称', '院（系）全称', '专业全称', '年级', '毕业时间', '联系电话', '邮箱', '团队分工']
    mt = b.doc.add_table(rows=4, cols=9)
    mt.style = 'Table Grid'
    for j, h in enumerate(headers):
        b.cell(mt.rows[0].cells[j], h, bold=True, size=Pt(8), center=True)
        b.shade(mt.rows[0].cells[j])
    for i, role in enumerate(['队长', '队员', '队员']):
        for j in range(8):
            b.cell(mt.rows[i + 1].cells[j], '【待填】', size=Pt(8), center=True)
        b.cell(mt.rows[i + 1].cells[8], role, size=Pt(8), center=True)

    b.text('团队指导教师信息（指导老师须与队长同校）', size=Pt(12), bold=True, space_before=8, space_after=4)
    gt = b.doc.add_table(rows=2, cols=6)
    gt.style = 'Table Grid'
    for j, h in enumerate(['姓名', '院（系）全称', '职称', '研究方向', '联系电话', '联系邮箱']):
        b.cell(gt.rows[0].cells[j], h, bold=True, size=Pt(9), center=True)
        b.shade(gt.rows[0].cells[j])
        b.cell(gt.rows[1].cells[j], '【待填写】', size=Pt(9), center=True)

    b.text('团队成员优势描述', size=Pt(12), bold=True, space_before=8, space_after=4)
    b.body(
        '本团队围绕「合同读不懂、风险看不见」这一校园与生活高频痛点组队，成员能力互补：'
        '（1）鸿蒙客户端：负责 ArkTS/ArkUI 主链路、应用接续/互动卡片/隐私防窥等系统能力接入与交互打磨；'
        '（2）后端与 AI：负责 Node.js 分析服务、大模型结构化输出、RAG 知识检索与法条引用校验；'
        '（3）产品与体验：负责信息架构（总览→风险列表→条款精读→追问）、样例合同与答辩演示预案。'
        '团队已完成可答辩版本：真实大模型分析主链路跑通，异常场景可降级演示，三项鸿蒙特性代码闭环。'
        '（请在提交前补充成员个人学术成果/项目经历与具体姓名分工。）'
    )
    b.page_break()

    # ========== 二、声明 ==========
    b.h('二、《法言白话》作品原创性声明', 1)
    b.body(
        '郑重声明：承诺本参赛队伍报名信息真实有效；呈交的参赛作品相关资料以及所完成的作品实物等相关成果，'
        '是本团队独立进行研究工作所取得的成果，除文中已经注明引用的内容外，本作品说明文档不包含任何其他个人或集体'
        '已经发表或撰写过的作品成果，不侵犯任何第三方的知识产权或其他权利。本声明的法律结果由本参赛队承担。'
    )
    b.text('参赛队员签名（团队全部成员）：______________________________', space_before=12, space_after=4)
    b.text('日期：____年____月____日', space_after=8)
    b.text('指导老师审核签名：______________________________', space_before=6, space_after=4)
    b.text('日期：____年____月____日', space_after=6)
    b.text(
        '（注：本页签名可以打印纸质文件后签名，提供扫描件；或者使用电子签名。不论哪种方式，需保证页面内容完整、字迹清晰，'
        '请与本项目创意书作为同一文档提交，否则项目创意书提交可能无效。）',
        size=Pt(9), color=RGBColor(0x55, 0x55, 0x55)
    )
    b.page_break()

    # ========== 三、创意 ==========
    b.h('三、创意描述', 1)
    b.text('一句话抽取关键创新点（不超过30字）：', space_after=3)
    b.text('签前看懂合同：摘要、风险分级、白话对照', size=Pt(14), bold=True,
           align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1D, 0x4E, 0xD8), space_before=4, space_after=2)
    b.text('（字数：19字）', size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x66, 0x66, 0x66), space_after=6)

    # 创意描述后直接接第四章，避免整页留白
    b.h('四、设计稿 / 技术方案', 1)

    b.h('4.1 产品定位与赛题契合', 2)
    b.body(
        '《法言白话》定位为基于 HarmonyOS 的「小而美」合同智能解读应用，赛题方向选择应用创新。'
        '作品面向租房、实习、培训、服务协议等校园与生活高频场景，将大模型能力与鸿蒙全场景能力结合：'
        '用 AI 解决「看不懂、找不准、怕踩坑」，用鸿蒙解决「跨端续看、桌面回看、隐私防窥」。'
        '已落地 ≥3 项鸿蒙特性：应用接续、互动卡片、隐私防窥，符合赛题对技术创新、体验新颖、场景实用与 AI 融合的要求。'
    )
    b.body(
        '产品边界明确：不做电子签署、律师咨询、合同自动改写、多人协作与法律意见书导出；'
        'OCR 为轻量提取与兜底，核心交付是大模型结构化解读与可理解的分层阅读体验。'
    )

    b.h('4.2 交互设计稿', 2)
    b.h('4.2.1 主交互流程', 3)
    b.body('主链路采用「分层阅读」：先结论、再重点、后精读，必要时追问，避免一次性抛出长文造成认知过载。')
    b.picture('fig1_interaction_flow.png')
    b.caption('图1 《法言白话》主交互流程')

    b.h('4.2.2 结果阅读信息架构', 3)
    b.body(
        'L1 总览给出整体摘要与签署前确认项；L2 用红/黄/绿快速定位；'
        'L3 条款页对照原文、白话、风险原因与修改建议；L4 在分析结果上下文中自然语言追问。'
    )
    b.picture('fig3_info_architecture.png')
    b.caption('图3 结果阅读信息架构（由浅入深）')

    # 4.3 文字技术路线（紧凑排版，不分多余页）
    write_section_43(b)

    b.h('4.4 技术实现方案', 2)
    b.body(
        '本节在 4.3 文字技术路线基础上，按模块说明「解决什么问题、用户看到什么、系统如何协作」。'
        '不展开源码清单与接口字段表，细节以仓库实现为准。'
    )

    b.h('4.4.1 合同导入与文本预处理', 3)
    b.body(
        '该模块解决「合同怎么进来」。用户可在导入页选择拍照、相册或本地文件，也可使用高低风险样例快速体验。'
        '对合同照片，客户端可先做识文，再把正文交给分析服务；对文本内容，则直接进入后续流程。'
        '服务端拿到正文后会做清洗与段落提示切分，为知识检索和大模型分析做准备。'
        '若文件格式尚不支持直接解析，系统会给出明确提示，而不是静默失败或假装读完。'
    )

    b.h('4.4.2 结构化风险分析', 3)
    b.body(
        '该模块解决「分析完给什么」。输出不是长篇评论文，而是可直接上屏的结构化结论：'
        '整体摘要告诉用户这份合同总体怎样；签署前清单告诉用户签字前先确认什么；'
        '红黄绿统计告诉用户重点在哪里；条款列表则提供原文、白话、原因与建议。'
        '模型被约束只能基于本合同原文解释，法条只能在有知识库依据时出现，避免「看起来很专业、其实不可追溯」。'
    )
    b.table([
        ['结果内容', '对用户的意义'],
        ['整体摘要', '先抓住总判断'],
        ['签署前确认项', '知道签字前该核对什么'],
        ['红 / 黄 / 绿统计', '快速分辨高风险与可保留点'],
        ['条款白话与建议', '看懂条款并知道可协商方向'],
        ['有据法条（可选）', '在有依据时增强可信度'],
    ])

    b.h('4.4.3 知识检索与法条校验', 3)
    b.body(
        '该模块解决「依据从哪来」。知识库提供风险规则、样例经验与法条摘录；'
        '分析服务先检索再生成，再在结果里核对法条是否属于本轮允许范围。'
        '没有检索到可用法条时，系统宁可留空，也不输出编造条文。'
        '知识库不可用时，分析仍可继续，只是不再给出具体法条引用，从而保证主链路可用。'
    )

    b.h('4.4.4 合同追问', 3)
    b.body(
        '该模块解决「看完结果还有问题怎么办」。用户可围绕本次合同继续提问，'
        '回答绑定已有分析结果与合同原文，并附带免责声明。'
        '它是精读之后的澄清工具，不是替代律师的开放咨询台。'
    )

    b.h('4.4.5 本地记录与任务状态', 3)
    b.body(
        '分析过程以任务状态推进，用户能感知「正在识别整理」还是「正在智能分析」。'
        '客户端会保存最近若干条分析记录，方便首页回看，也为桌面卡片提供内容。'
        '设置项如服务地址、隐私开关、演示模式等保存在本地，便于真机联调与答辩切换。'
        '服务端当前以内存保存任务，适合演示；这与长期生产存储不是同一回事。'
    )

    b.h('4.4.6 服务协作方式', 3)
    b.body(
        '客户端与分析服务之间采用「先建任务、再查进度、最后取结果」的协作方式，'
        '并支持基于本次结果的追问。健康检查用于确认服务是否就绪。'
        '出错时返回可读原因，页面展示重试或替代路径，避免用户卡在无提示等待中。'
    )

    b.h('4.4.7 界面展示与反馈', 3)
    b.body(
        '页面覆盖首页、导入、分析中、结果总览、条款对照、合同问答与设置。'
        '统一的加载态、空态与错误态，让等待和失败都可理解、可操作。'
        '真机联调时可在设置中修改服务地址；异常时可切换演示样例，保证讲解不中断。'
    )

    b.h('4.4.8 鸿蒙全场景能力', 3)
    b.body(
        '应用接续支持在结果阅读场景下跨设备继续查看同一份分析；'
        '互动卡片支持在桌面回看最近合同摘要并进入结果；'
        '隐私防窥支持在敏感阅读页降低截录屏风险。'
        '三项能力都服务于「签前读合同」这件事，而不是为了凑特性数量。'
    )

    b.h('4.4.9 报告导出与演示兜底', 3)
    b.body(
        '用户可将分析结果导出为便于留存的报告，内容包含摘要与风险明细，并再次声明不构成正式法律意见。'
        '答辩场景下，可通过本地样例与弱网、超时、服务失败、降级等模式，主动演示异常反馈，'
        '不依赖现场网络偶然性。'
    )

    b.h('4.5 系统实现效果', 2)
    b.body(
        '当前版本已形成可答辩闭环：导入后能完成分析，结果页能看懂摘要与风险，条款页能对照白话，'
        '必要时可追问；知识库可用时提供有据法条，不可用时仍能给出结构化解读；'
        '首页可回看最近记录，报告可导出备忘；鸿蒙三项特性具备可讲解的实现链路；'
        '异常路径可主动演示。跨设备接续建议在同账号设备上彩排确认。'
    )
    b.table([
        ['效果项', '用户可感知结果'],
        ['主链路', '导入到对照/追问可走通'],
        ['智能解读', '摘要、红黄绿、白话与建议可见'],
        ['依据约束', '无法条依据时不编造条文'],
        ['回看与导出', '最近记录与报告备忘可用'],
        ['鸿蒙能力', '接续、卡片、防窥可讲述'],
        ['现场兜底', '样例与异常模式可切换'],
    ])

    b.h('4.6 运行与使用说明', 2)
    b.body(
        '运行时先启动知识库并完成索引，再启动分析服务并配置模型密钥，最后用 DevEco Studio 运行客户端。'
        '真机联调请使用电脑局域网地址，不要使用仅对本机有效的回环地址。'
        '仓库按客户端、分析服务、知识库与文档分目录组织，便于分工与核对。'
        '具体命令与环境变量以 README 及知识库说明为准。'
    )

    # ========== 五、介绍 ==========
    b.h('五、介绍文档', 1)
    b.text('（不超过800字；以下为正式提交正文）', size=Pt(9), color=RGBColor(0x66, 0x66, 0x66), space_after=8)
    intro = (
        '租房、实习、培训与服务协议里，条款冗长且充满专业表述，普通人常读不完、看不懂，风险容易藏在字缝里，签完才发现被动。'
        '《法言白话》面向非法律专业用户，基于 HarmonyOS 打造「签合同前先看懂」的智能解读应用：导入合同后，先给出整体摘要与签署前确认项，'
        '再用红、黄、绿风险分级快速定位重点，并在条款页对照展示原文、白话解释、风险原因与修改建议；必要时可围绕本次分析结果继续自然语言追问。'
        '作品明确不替代律师，目标是降低理解门槛、突出必须确认的关键点。'
        '实现上采用端—云—知识库协同：ArkTS 客户端负责导入、状态轮询与分层展示；Node 分析服务完成文本清洗、条款切分与大模型结构化输出；'
        'Python RAG 检索风险规则与法条摘录，经引用校验后注入提示词，无知识库依据不输出法条，知识库不可用时可降级为纯 Prompt，保证主链路可用。'
        '客户端不直连大模型，模型调用统一收敛于服务端，便于密钥管理与现场演示兜底。'
        '鸿蒙侧已落地应用接续、互动卡片与隐私防窥：分别支撑手机分析后平板/PC 续看、桌面卡片回看最近摘要，以及敏感结果页限制截屏录屏，贴合全场景与安全诉求。'
        '应用前景上，高校与青年群体在租房、实习、就业签约中持续存在「先读懂再签字」的刚需，轻量解读工具具备清晰使用场景与扩散空间；'
        '首版聚焦解读主链路的稳定可用与可讲清楚，后续可按场景扩展合同类型与端侧体验。'
    )
    b.body(intro)
    b.text(f'【介绍文档正文字数：{len(intro)} 字】', size=Pt(9), align=WD_ALIGN_PARAGRAPH.RIGHT,
           color=RGBColor(0x66, 0x66, 0x66))

    b.h('5.1 补充说明（不计入 800 字正文）', 2)
    b.body(
        '首版边界：不做电子签署、律师咨询、合同自动改写、多人协作与法律意见书导出；'
        'OCR 为轻量提取+兜底，核心交付是大模型结构化解读。答辩可演示高低风险样例及异常模式。'
    )

    b.save()


if __name__ == '__main__':
    build()
